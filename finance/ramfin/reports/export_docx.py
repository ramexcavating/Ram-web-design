"""Weekly report as a Word document, matching the existing YYMMDD_RAM_Weekly_Managers_Report.docx habit.
Converts the markdown the report already produces: headings, bullets, bold, pipe tables."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def _runs(par, text: str) -> None:
    for i, part in enumerate(re.split(r"(\*\*[^*]+\*\*)", text)):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = par.add_run(part[2:-2]); r.bold = True
        else:
            par.add_run(part)


def markdown_to_docx(md: str, path: str | Path, title_suffix: str = "") -> Path:
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        ln = lines[i]
        if ln.startswith("# "):
            doc.add_heading(ln[2:] + title_suffix, 0)
        elif ln.startswith("## "):
            doc.add_heading(ln[3:], 1)
        elif ln.startswith("### "):
            doc.add_heading(ln[4:], 2)
        elif ln.startswith("|"):
            block = []
            while i < len(lines) and lines[i].startswith("|"):
                block.append(lines[i]); i += 1
            rows = [[c.strip() for c in r.strip().strip("|").split("|")] for r in block if not re.match(r"^\|[\s:|-]+\|$", r)]
            if rows:
                t = doc.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Light Grid Accent 1"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row[:len(rows[0])]):
                        c = t.cell(ri, ci); c.text = ""
                        p = c.paragraphs[0]; _runs(p, cell)
                        if ri == 0:
                            for r in p.runs: r.bold = True
                        elif re.match(r"^-?[\d,]+(\.\d+)?%?$", cell):
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        if cell == "BREACH":
                            for r in p.runs: r.font.color.rgb = RGBColor(0xC0, 0, 0); r.bold = True
                        elif cell == "TIGHT":
                            for r in p.runs: r.font.color.rgb = RGBColor(0xA3, 0x6A, 0)
            continue
        elif ln.startswith("- "):
            _runs(doc.add_paragraph(style="List Bullet"), ln[2:])
        elif ln.strip():
            _runs(doc.add_paragraph(), ln)
        i += 1
    p = Path(path); p.parent.mkdir(parents=True, exist_ok=True)
    doc.save(p)
    return p
