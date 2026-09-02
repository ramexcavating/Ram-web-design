"""Equipment fold-in, Word export, state sync and the delegated mailbox path."""
from datetime import date
from pathlib import Path

from docx import Document

from ramfin import db, pipeline
from ramfin.extract.extractor import FakeExtractor
from ramfin.filer import LocalFiler
from ramfin.ledger import intake, jobcost
from ramfin.models import Extraction
from ramfin.reports.export_docx import markdown_to_docx
from ramfin.sources.graph import GraphClient
from tests.fixtures import extractions as fx


def test_equipment_unit_detection():
    assert intake.equipment_unit(None, Extraction.from_dict(fx.INVOICE_BRANDT)) == "EX-03"
    assert intake.equipment_unit(None, Extraction.from_dict({**fx.RECEIPT_NO_JOB, "handwritten_equipment_id": "dt 02"})) == "DT-02"
    assert intake.equipment_unit(None, Extraction.from_dict(fx.RECEIPT_HOME_HARDWARE)) is None


def test_equipment_copy_filed_and_costed(conn, settings, tmp_path):
    drop = tmp_path / "drop"; drop.mkdir()
    (drop / "brandt.pdf").write_bytes(b"x")
    (drop / "fuel.jpg").write_bytes(b"y")
    from ramfin.sources.folder import scan_local_folder
    scan_local_folder(conn, settings, drop)
    ex = FakeExtractor({"brandt.pdf": fx.INVOICE_BRANDT,
                        "fuel.jpg": {**fx.RECEIPT_NO_JOB, "handwritten_job_no": "240617", "handwritten_cost_code": "02-300", "handwritten_equipment_id": "EX-03"}})
    pipeline.process_new_documents(conn, settings, ex, LocalFiler(tmp_path / "sp"))
    copies = [p.as_posix() for p in (tmp_path / "sp").rglob("*") if p.is_file() and "01_SERVICE_RECORDS" in p.as_posix()]
    assert len(copies) == 2 and all("05_EQUIPMENT/01_FLEET/EX-03/01_SERVICE_RECORDS/" in c for c in copies)
    ec = jobcost.equipment_cost(conn, settings)
    assert ec["EX-03"]["repairs"] == 3120.0 and round(ec["EX-03"]["fuel"], 2) == round(212.40 - 10.11, 2)
    assert conn.execute("SELECT sharepoint_folder FROM equipment WHERE unit_id='EX-03'").fetchone()["sharepoint_folder"] == "05_EQUIPMENT/01_FLEET/EX-03"


def test_markdown_to_docx(tmp_path):
    md = "# Report\n\n**Headline here.**\n\n## Table\n\n| Week | Closing | Status |\n|---|---:|---|\n| 2026-08-31 | -35,000 | BREACH |\n\n- a bullet\n"
    p = markdown_to_docx(md, tmp_path / "r.docx")
    d = Document(p)
    texts = [x.text for x in d.paragraphs]
    assert "Report" in texts and any("Headline here." in t for t in texts) and "a bullet" in texts
    assert d.tables[0].cell(1, 2).text == "BREACH"


def test_weekly_writes_docx(conn, settings, tmp_path):
    db.insert(conn, "bank_balances", dict(as_of="2026-08-31", account_key="rbc_chq", balance=1000.0, source="manual", created_at=db.now_iso()))
    out = pipeline.weekly(conn, settings, tmp_path, as_of=date(2026, 8, 31))
    assert Path(out["docx"]).exists() and Path(out["docx"]).name == "260831_RAM_Weekly_Managers_Report.docx"


def test_me_mailbox_urls():
    assert GraphClient._mbx("me") == "/me"
    assert GraphClient._mbx("accounts@ramexcavating.ca") == "/users/accounts@ramexcavating.ca"


def test_state_sync_roundtrip(conn, settings, tmp_path, monkeypatch):
    from ramfin import state_sync
    store: dict[str, bytes] = {}

    class FakeGraph:
        def upload_replace(self, drive, folder, name, data):
            store[f"{folder}/{name}"] = data; return f"https://sp/{folder}/{name}"
        def download_path(self, drive, path):
            return store.get(path)

    real = db.connect(settings.db_path)
    real.execute("INSERT INTO jobs(job_no, name) VALUES('999999','roundtrip')"); real.commit(); real.close()
    (settings.data_dir / state_sync.TOKEN_NAME).write_text("{}")
    pushed = state_sync.push(FakeGraph(), "drive", settings)
    assert set(pushed) == {state_sync.DB_NAME, state_sync.TOKEN_NAME}
    settings.db_path.unlink()
    pulled = state_sync.pull(FakeGraph(), "drive", settings)
    assert pulled[state_sync.DB_NAME] and db.connect(settings.db_path).execute("SELECT name FROM jobs WHERE job_no='999999'").fetchone()["name"] == "roundtrip"
